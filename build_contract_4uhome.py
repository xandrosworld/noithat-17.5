from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\Admin\Desktop\KHACHHANG\noithat-17.5")
OUT = ROOT / "Hop-dong-mau-4UHome-UIUX-Website-v2-seo-update.docx"
LOGO = ROOT / "FILE BÀN GIAO LOGO" / "LOGO-01 PNG.png"

GOLD = "D9B66B"
LIGHT = "F7F1E4"
PALE = "FBF8F1"
BORDER = "D9C99B"
FONT = "Calibri"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def border_cell(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, size=9.5, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_style(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            border_cell(cell)
            cell.margin_top = Cm(0.08)
            cell.margin_bottom = Cm(0.08)
            cell.margin_left = Cm(0.12)
            cell.margin_right = Cm(0.12)
        if header and row_index == 0:
            for cell in row.cells:
                shade(cell, GOLD)


def add_heading(doc, num, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    r = p.add_run(f"{num:02d}  {title}")
    r.bold = True


def add_note(doc, text):
    t = doc.add_table(rows=1, cols=1)
    table_style(t, header=False)
    shade(t.cell(0, 0), PALE)
    set_cell_text(t.cell(0, 0), text, size=9.2)
    doc.add_paragraph("")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.2)
        p.add_run("• ").bold = True
        p.add_run(item)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = FONT
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(153, 104, 26)

    cover = doc.add_table(rows=1, cols=1)
    table_style(cover, header=False)
    shade(cover.cell(0, 0), PALE)
    p = cover.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Cm(5.0))
    p = cover.cell(0, 0).add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HỢP ĐỒNG DỊCH VỤ THIẾT KẾ UI/UX & PHÁT TRIỂN WEBSITE")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(153, 104, 26)
    p = cover.cell(0, 0).add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DỰ ÁN WEBSITE NỘI THẤT 4U HOME")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(51, 51, 51)
    p = cover.cell(0, 0).add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hợp đồng mẫu + Phụ lục phạm vi tính năng + Biên bản nghiệm thu")
    r.font.size = Pt(10)
    r.italic = True

    doc.add_paragraph("")
    summary = doc.add_table(rows=2, cols=4)
    table_style(summary, header=False)
    for i, (k, v) in enumerate(
        [
            ("Thương hiệu", "4U HOME"),
            ("Ngày lập mẫu", "17/05/2026"),
            ("Trang tham chiếu", "hungphuthinh.vn"),
            ("Tình trạng", "Thông tin pháp lý để trống"),
        ]
    ):
        shade(summary.cell(0, i), GOLD)
        set_cell_text(summary.cell(0, i), k, bold=True, size=8.8, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(summary.cell(1, i), v, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    add_heading(doc, 1, "THÔNG TIN HỢP ĐỒNG & CÁC BÊN")
    info_rows = [
        "Tên tổ chức/cá nhân",
        "Đại diện",
        "Chức vụ",
        "Mã số thuế/CCCD",
        "Số điện thoại",
        "Email",
        "Địa chỉ",
        "Tài khoản thanh toán",
        "Người trao đổi dự án",
        "Tên dự án",
        "Website tham chiếu",
        "Hình thức ký",
    ]
    t = doc.add_table(rows=len(info_rows) + 1, cols=3)
    table_style(t)
    for i, h in enumerate(["MỤC", "BÊN A - BÊN THUÊ DỊCH VỤ", "BÊN B - BÊN CUNG CẤP DỊCH VỤ"]):
        set_cell_text(t.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row, label in enumerate(info_rows, start=1):
        set_cell_text(t.cell(row, 0), label, bold=True, size=8.8)
        a = "Website nội thất 4U HOME" if label == "Tên dự án" else (
            "https://hungphuthinh.vn/" if label == "Website tham chiếu" else (
                "Ký online/bản mềm hoặc bản giấy theo thống nhất hai bên" if label == "Hình thức ký" else "................................................................................"
            )
        )
        b = "Thiết kế UI/UX, phát triển web demo thật, website chính và hệ landing page/site vệ tinh" if label == "Tên dự án" else (
            "https://hungphuthinh.vn/" if label == "Website tham chiếu" else (
                "Ký online/bản mềm hoặc bản giấy theo thống nhất hai bên" if label == "Hình thức ký" else "................................................................................"
            )
        )
        set_cell_text(t.cell(row, 1), a, size=8.6)
        set_cell_text(t.cell(row, 2), b, size=8.6)

    add_heading(doc, 2, "CĂN CỨ KÝ KẾT")
    add_bullets(
        doc,
        [
            "Bộ luật Dân sự năm 2015, Luật Thương mại năm 2005 và các quy định pháp luật Việt Nam có liên quan.",
            "Nhu cầu của Bên A về việc thiết kế giao diện UI/UX và phát triển website cho thương hiệu 4U HOME trong lĩnh vực nội thất/nhà ở.",
            "Trao đổi ngày 17/05/2026: Bên A mong muốn làm web demo thật để duyệt giao diện, chức năng tương tự website tham chiếu, đồng thời bổ sung landing page/site vệ tinh, private blog network và subdomain để kéo traffic về trang mẹ.",
            "Năng lực thiết kế, phát triển, triển khai và bàn giao website của Bên B.",
            "Các phụ lục phạm vi tính năng, timeline, báo giá và biên bản nghiệm thu là một phần không tách rời của hợp đồng.",
        ],
    )

    add_heading(doc, 3, "KẾT QUẢ NGHIÊN CỨU WEBSITE THAM CHIẾU")
    research = [
        ("Header, menu, search", "Logo, menu đa cấp, ô tìm kiếm, banner nhanh, mobile menu.", "Header responsive, menu danh mục nội thất/dự án/báo giá/tin tức, tìm kiếm nội dung."),
        ("Trang chủ", "Hero slider, cụm dịch vụ, giới thiệu, video, thống kê, công trình thực tế, cam kết, mẫu thiết kế, đội ngũ, đối tác, footer SEO.", "Trang chủ 4U HOME có hero theo nhận diện logo, CTA tư vấn, dự án/mẫu nội thất nổi bật, cam kết, đối tác, social."),
        ("Báo giá/dịch vụ", "Các trang báo giá xây nhà trọn gói, phần thô, thiết kế; nội dung dài chuẩn SEO.", "Trang báo giá/dịch vụ nội thất, form lấy nhu cầu, bảng gói dịch vụ để khách dễ duyệt."),
        ("Dự án/công trình", "Danh sách dự án, phân loại công trình đang thi công/đã bàn giao, trang chi tiết có ảnh và nội dung.", "Module dự án nội thất, bộ lọc loại công trình/phong cách/khu vực, chi tiết dự án và gallery."),
        ("Thiết kế nhà đẹp/nội thất", "Danh mục biệt thự, nhà phố, tầng; nội thất theo phòng: khách, bếp, ngủ, trẻ em, restroom, cảnh quan.", "Danh mục sản phẩm/dịch vụ nội thất theo không gian, phong cách và nhu cầu sử dụng."),
        ("Tin tức SEO", "Kho bài viết lớn về chi phí, phong thủy, kinh nghiệm, khu vực, mẫu nhà.", "Blog SEO + landing page theo từ khóa/dịch vụ/khu vực để phục vụ kéo traffic."),
        ("Liên hệ/CTA", "Form liên hệ, hotline, Zalo, social, bản đồ/khu vực phục vụ.", "Form tư vấn, nút gọi/Zalo cố định, cấu hình thông tin liên hệ và tracking chuyển đổi."),
        ("SEO kỹ thuật", "Sitemap, canonical, analytics/tag manager, schema, link SEO footer.", "SEO cơ bản, sitemap, robots, schema Organization/LocalBusiness, analytics, landing page indexable."),
    ]
    t = doc.add_table(rows=len(research) + 1, cols=3)
    table_style(t)
    for i, h in enumerate(["NHÓM MÀN HÌNH/CHỨC NĂNG", "GHI NHẬN TỪ HUNGPHUTHINH.VN", "ĐƯA VÀO DỰ ÁN 4U HOME"]):
        set_cell_text(t.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(research, start=1):
        for c, val in enumerate(row):
            set_cell_text(t.cell(r, c), val, size=8.2)

    add_heading(doc, 4, "PHẠM VI CÔNG VIỆC")
    scope = [
        ("UI/UX web demo thật", "Thiết kế và dựng giao diện demo responsive: desktop, tablet, mobile; dùng logo/màu sắc 4U HOME; ưu tiên trải nghiệm xem mẫu nội thất và gửi yêu cầu tư vấn.", "Không chỉ file Figma; có bản web demo để khách duyệt."),
        ("Website chính", "Trang chủ, giới thiệu, dịch vụ/báo giá, dự án, chi tiết dự án, danh mục nội thất, bài viết/tin tức, liên hệ, tìm kiếm.", "Chức năng tương tự website tham chiếu, điều chỉnh cho ngành nội thất."),
        ("Quản trị nội dung/CMS", "Admin quản lý banner, menu, trang tĩnh, dịch vụ/báo giá, dự án, gallery, blog, landing page SEO, form liên hệ, lead, thông tin footer/social và cấu hình xuất bản.", "Bàn giao tài khoản admin; chi tiết trường CMS phục vụ SEO/performance ở bảng bên dưới."),
        ("Landing page/site vệ tinh", "Tạo cấu trúc landing page theo dịch vụ/khu vực/từ khóa; có template nhân bản nhanh và liên kết về trang mẹ.", "Phục vụ SEO/PBN theo yêu cầu, triển khai theo cấu hình domain/subdomain thực tế."),
        ("Subdomain", "Hỗ trợ cấu hình subdomain cho site vệ tinh/landing page nếu Bên A cung cấp quyền DNS/hosting.", "Số lượng subdomain cụ thể chốt ở phụ lục/báo giá."),
        ("SEO & tracking", "SEO on-page cơ bản, URL thân thiện, sitemap, robots, metadata, schema cơ bản, canonical, redirect 301/302 khi cập nhật URL, analytics/tag manager, Meta Pixel/Facebook Pixel và event cơ bản theo mã tracking Bên A cung cấp.", "Không cam kết thứ hạng SEO, hiệu quả Google Ads/Facebook Ads hoặc doanh thu; chỉ bàn giao nền tảng kỹ thuật và tracking cơ bản."),
        ("Quản lý thẻ client-side", "Bổ sung khu vực cấu hình các thẻ HTML/script phía client như head/body scripts, mã xác minh, pixel, conversion tag hoặc mã đo lường hợp lệ.", "Bên A chịu trách nhiệm cung cấp mã đúng từ nền tảng quảng cáo/đo lường; Bên B hỗ trợ gắn và kiểm tra hiển thị ở mức cơ bản."),
        ("Tối ưu & bàn giao", "Responsive, tối ưu tốc độ cơ bản, hướng dẫn quản trị, bàn giao mã nguồn/tài khoản theo điều kiện thanh toán.", "Hosting, domain, plugin/trả phí bên thứ ba do Bên A chi trả."),
    ]
    t = doc.add_table(rows=len(scope) + 1, cols=3)
    table_style(t)
    for i, h in enumerate(["NHÓM", "BAO GỒM TRONG GÓI MẪU", "GHI CHÚ PHẠM VI"]):
        set_cell_text(t.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(scope, start=1):
        for c, val in enumerate(row):
            set_cell_text(t.cell(r, c), val, size=8.2)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Chi tiết chức năng CMS phục vụ SEO/performance")
    r.bold = True
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("A36A16")
    cms_detail = [
        ("Trang tĩnh & menu", "Tạo/sửa trang giới thiệu, dịch vụ, báo giá; sắp xếp menu header/footer; bật/tắt hiển thị; xem trước trước khi xuất bản.", "Slug, SEO title, meta description, canonical, robots index/follow, breadcrumb schema."),
        ("Banner trang chủ", "Quản lý slide hero, ảnh desktop/mobile, tiêu đề, mô tả, CTA, thứ tự, lịch hiển thị và trạng thái chiến dịch.", "Alt ảnh, lazy-load, link CTA có UTM, event click_banner cho GA4/Pixel."),
        ("Dự án/Gallery", "Thêm dự án, phân loại công trình/phong cách/khu vực, nhập thông số, câu chuyện thiết kế, vật liệu và album ảnh.", "Project schema, alt/caption ảnh, internal link tới dịch vụ liên quan, URL dự án thân thiện."),
        ("Blog tin tức", "Quản lý danh mục, bài viết, tác giả, ngày xuất bản/cập nhật, mục lục, FAQ, bài liên quan và trạng thái duyệt.", "Article schema, FAQ schema, OG image, focus keyword, internal links, sitemap blog."),
        ("Landing page SEO", "Nhân bản template theo dịch vụ/khu vực/từ khóa, thêm block nội dung, form CTA, FAQ, bảng giá và liên kết về trang mẹ.", "Index/noindex, canonical, sitemap priority, keyword mapping, conversion event theo từng landing."),
        ("Lead & form liên hệ", "Lưu lead từ form, phân loại nguồn, trạng thái tư vấn, ghi chú, xuất CSV và đánh dấu lead đã xử lý.", "Event lead_submit, phone_click, zalo_click; đo nguồn UTM và trang chuyển đổi."),
        ("SEO kỹ thuật", "Quản lý metadata toàn site/từng trang, schema, sitemap, robots, redirect, scripts đo lường và mã xác minh.", "Sitemap.xml, robots.txt, JSON-LD, canonical tags, GTM/GA4/Meta Pixel."),
    ]
    t = doc.add_table(rows=len(cms_detail) + 1, cols=3)
    table_style(t)
    for i, h in enumerate(["NHÓM CMS", "CHỨC NĂNG QUẢN TRỊ CỤ THỂ", "TRƯỜNG/KẾT QUẢ SEO - TRACKING"]):
        set_cell_text(t.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_idx, row in enumerate(cms_detail, start=1):
        for col_idx, val in enumerate(row):
            set_cell_text(t.cell(row_idx, col_idx), val, size=7.6)
    add_note(doc, "Lưu ý bảo vệ phạm vi: Các tính năng phát sinh như CRM nâng cao, chatbot AI, automation SEO hàng loạt, viết nội dung số lượng lớn, import dữ liệu lớn, app mobile, thiết kế brand guideline đầy đủ hoặc bảo trì dài hạn chỉ thực hiện khi có phụ lục/báo giá bổ sung.")

    add_heading(doc, 5, "GIÁ TRỊ HỢP ĐỒNG & THANH TOÁN")
    pay = [
        ("5.1", "Gói website chính 4U HOME: UI/UX web demo thật, frontend responsive, CMS quản trị nội dung chi tiết, trang chủ, dịch vụ/báo giá, dự án, chi tiết dự án, danh mục nội thất, blog/tin tức, liên hệ, SEO cơ bản", "18.000.000 VNĐ"),
        ("5.2", "Gói landing page/site vệ tinh: template landing page theo từ khóa/dịch vụ/khu vực, cấu trúc liên kết về trang mẹ, hỗ trợ cấu hình subdomain khi Bên A cung cấp quyền DNS/hosting", "5.000.000 VNĐ"),
        ("5.3", "Tổng giá trị hợp đồng đề xuất, đã tham chiếu gói mẫu 20.000.000 VNĐ và điều chỉnh theo scope hiện tại", "23.000.000 VNĐ"),
        ("5.4", "Đợt 1: 35% trong vòng 01 ngày kể từ ngày ký hợp đồng", "8.050.000 VNĐ"),
        ("5.5", "Đợt 2: 55% trong vòng 03 ngày kể từ khi duyệt demo chính/admin và trước khi bàn giao mã nguồn", "12.650.000 VNĐ"),
        ("5.6", "Đợt 3: 10% sau khi nghiệm thu và vận hành ổn định 30 ngày trong phạm vi bảo hành", "2.300.000 VNĐ"),
        ("5.7", "Phí hosting, tên miền, SSL, email, công cụ AI, SMS/Zalo, plugin trả phí, quảng cáo hoặc dịch vụ bên thứ ba", "Bên A thanh toán trực tiếp khi phát sinh"),
        ("5.8", "Hình thức thanh toán", "Chuyển khoản/tiền mặt theo thông tin hai bên xác nhận"),
        ("5.9", "Điều kiện triển khai", "Bên B bắt đầu sau khi hai bên ký và Bên A thanh toán đợt 1"),
    ]
    t = doc.add_table(rows=len(pay) + 1, cols=3)
    table_style(t)
    for i, h in enumerate(["MỤC", "NỘI DUNG", "GIÁ TRỊ/THỜI HẠN"]):
        set_cell_text(t.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(pay, start=1):
        for c, val in enumerate(row):
            set_cell_text(t.cell(r, c), val, size=8.4)

    add_heading(doc, 6, "THỜI GIAN THỰC HIỆN DỰ KIẾN")
    timeline = [
        ("Ngày 1-2", "Chốt sitemap, cấu trúc màn hình, style UI theo logo 4U HOME và trang tham chiếu."),
        ("Ngày 3-5", "Dựng web demo: trang chủ, dịch vụ/báo giá, dự án, danh mục nội thất, tin tức, liên hệ."),
        ("Ngày 6-8", "Xây admin CMS quản trị banner, dự án/gallery, blog/landing page, SEO fields, form liên hệ, lead, tìm kiếm, CTA Zalo/phone."),
        ("Ngày 9-11", "Triển khai landing page/site vệ tinh, template nhân bản, cấu hình SEO cơ bản và subdomain nếu đã có quyền truy cập."),
        ("Ngày 12-14", "Kiểm thử responsive, tối ưu cơ bản, chỉnh theo phản hồi, demo nghiệm thu."),
        ("Sau nghiệm thu", "Bàn giao mã nguồn/tài khoản/admin/hướng dẫn sử dụng theo điều kiện thanh toán."),
    ]
    t = doc.add_table(rows=len(timeline) + 1, cols=2)
    table_style(t)
    for i, h in enumerate(["MỐC", "CÔNG VIỆC DỰ KIẾN"]):
        set_cell_text(t.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(timeline, start=1):
        for c, val in enumerate(row):
            set_cell_text(t.cell(r, c), val, size=8.5)
    add_note(doc, "Timeline chỉ là mẫu tham khảo và có thể thay đổi tùy tốc độ phản hồi, số lượng nội dung, số lượng landing page/subdomain và dữ liệu do Bên A cung cấp.")

    add_heading(doc, 7, "QUYỀN, NGHĨA VỤ & BÀN GIAO")
    terms = [
        ("Quyền của Bên A", "Kiểm tra tiến độ, góp ý hợp lý, nghiệm thu hạng mục, nhận bàn giao sản phẩm trong phạm vi đã thanh toán."),
        ("Nghĩa vụ Bên A", "Cung cấp logo, nội dung, hình ảnh, tài khoản hosting/tên miền/DNS, phản hồi đúng thời hạn, thanh toán đúng tiến độ và chịu trách nhiệm pháp lý về nội dung cung cấp."),
        ("Quyền của Bên B", "Tạm dừng triển khai khi chậm thanh toán hoặc thiếu dữ liệu; từ chối yêu cầu ngoài phạm vi khi chưa có phụ lục bổ sung."),
        ("Nghĩa vụ Bên B", "Thiết kế, phát triển, kiểm thử và bàn giao đúng phạm vi; hỗ trợ cấu hình hosting/tên miền/subdomain trong khả năng kỹ thuật khi Bên A cung cấp quyền truy cập."),
        ("Sở hữu trí tuệ", "Sau khi Bên A thanh toán đầy đủ, Bên A được quyền sử dụng mã nguồn và sản phẩm đã bàn giao trong phạm vi hợp đồng."),
        ("Bảo mật", "Hai bên bảo mật thông tin nội bộ, tài khoản truy cập, dữ liệu dự án và không tiết lộ cho bên thứ ba nếu chưa được đồng ý."),
    ]
    t = doc.add_table(rows=len(terms) + 1, cols=2)
    table_style(t)
    for i, h in enumerate(["ĐIỀU KHOẢN", "NỘI DUNG TÓM TẮT"]):
        set_cell_text(t.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(terms, start=1):
        for c, val in enumerate(row):
            set_cell_text(t.cell(r, c), val, size=8.4)

    add_heading(doc, 8, "BẢO HÀNH, THAY ĐỔI PHẠM VI & TRANH CHẤP")
    add_bullets(
        doc,
        [
            "Bảo hành lỗi lập trình/cấu hình do Bên B gây ra trong ........ ngày kể từ ngày nghiệm thu. Yêu cầu thêm chức năng mới không thuộc bảo hành.",
            "Mọi yêu cầu ngoài phạm vi đã chốt cần có xác nhận chi phí/thời gian bằng phụ lục, báo giá hoặc tin nhắn xác nhận giữa hai bên.",
            "Bên thứ ba: hosting, tên miền, SSL, email, công cụ AI, plugin, dữ liệu trả phí, chi phí quảng cáo và phí dịch vụ nền tảng do Bên A chi trả.",
            "SEO/PBN: Bên B bàn giao nền tảng kỹ thuật và cấu trúc landing page/site vệ tinh; không cam kết thứ hạng, traffic hoặc doanh thu nếu không có hợp đồng SEO riêng.",
            "Tranh chấp: Hai bên ưu tiên thương lượng thiện chí. Nếu không thống nhất, áp dụng pháp luật Việt Nam và cơ quan có thẩm quyền theo thỏa thuận của hai bên.",
        ],
    )

    add_heading(doc, 9, "PHỤ LỤC 1 - CHECKLIST TÍNH NĂNG ĐÃ BAO GỒM")
    features = [
        ("Trang chủ theo nhận diện 4U HOME", "Header/menu đa cấp", "Hero banner/slider", "Search nội dung"),
        ("Trang giới thiệu", "Trang dịch vụ/báo giá", "Danh mục dự án", "Chi tiết dự án + gallery"),
        ("Danh mục nội thất theo không gian", "Danh mục theo phong cách", "Blog/tin tức SEO", "Trang liên hệ + form"),
        ("CTA gọi nhanh/Zalo", "Footer SEO + social", "Responsive mobile/tablet/desktop", "Admin quản lý nội dung"),
        ("Quản lý banner/menu", "Quản lý dự án/gallery", "Quản lý bài viết/danh mục", "Quản lý form liên hệ/lead"),
        ("CMS phân quyền admin/SEO", "Preview trước xuất bản", "Trạng thái index/noindex", "Theo dõi nguồn UTM"),
        ("Landing page theo từ khóa", "Template site vệ tinh", "Liên kết về trang mẹ", "Hỗ trợ subdomain"),
        ("Alt/caption ảnh từng hình", "FAQ/schema từng bài", "Mục lục/internal link", "Sitemap cập nhật khi publish"),
        ("SEO title/description", "URL thân thiện", "Sitemap/robots", "Schema cơ bản"),
        ("Canonical", "Redirect 301/302", "Meta Pixel/Facebook Pixel", "Analytics/Tag Manager"),
        ("Quản lý thẻ HTML client-side", "Tối ưu tốc độ cơ bản", "Hướng dẫn sử dụng admin", "Bàn giao mã nguồn theo thanh toán"),
    ]
    t = doc.add_table(rows=len(features) + 1, cols=4)
    table_style(t)
    for c in range(4):
        set_cell_text(t.cell(0, c), "TÍNH NĂNG", bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(features, start=1):
        for c, val in enumerate(row):
            set_cell_text(t.cell(r, c), "✓ " + val, size=8.0)

    add_heading(doc, 10, "PHỤ LỤC 2 - HẠNG MỤC MUA THÊM / CHỐT SAU")
    addons = [
        ("Viết nội dung SEO hàng loạt", "Sản xuất bài viết/landing page theo bộ từ khóa và khu vực, chưa gồm trong gói 23tr.", "300.000 - 700.000 VNĐ/bài", "Nên mua khi chạy SEO"),
        ("Thêm landing page ngoài gói", "Nhân bản thêm landing page riêng theo dịch vụ/khu vực/từ khóa sau số lượng đã chốt.", "300.000 - 800.000 VNĐ/trang", "Tùy chiến dịch"),
        ("Thêm site vệ tinh/subdomain", "Cấu hình thêm site vệ tinh hoặc subdomain ngoài phạm vi ban đầu.", "1.000.000 - 2.000.000 VNĐ/site", "Tùy chọn"),
        ("Automation PBN nâng cao", "Tự động tạo/nhân bản landing page, lịch đăng bài, internal link theo rule.", "+5.000.000 - 10.000.000 VNĐ", "Nâng cấp"),
        ("CRM khách hàng", "Quản lý lead, trạng thái tư vấn, phân công nhân sự, lịch hẹn.", "+5.000.000 - 8.000.000 VNĐ", "Nên mua"),
        ("Chatbot AI tư vấn", "Chatbot thu thập nhu cầu, tư vấn mẫu nội thất/gói dịch vụ cơ bản.", "+5.000.000 - 12.000.000 VNĐ", "Tùy chọn"),
        ("Thư viện sản phẩm/nội thất", "Quản lý sản phẩm, vật liệu, bộ sưu tập, báo giá theo lựa chọn.", "+5.000.000 - 10.000.000 VNĐ", "Nên mua"),
        ("Đa ngôn ngữ", "Tiếng Việt/Anh hoặc ngôn ngữ khác.", "+3.000.000 - 6.000.000 VNĐ", "Tùy chọn"),
        ("Tích hợp thanh toán", "Đặt cọc online/QR/cổng thanh toán.", "+3.000.000 - 7.000.000 VNĐ", "Tùy chọn"),
        ("Bảo trì tháng", "Theo dõi lỗi, backup, cập nhật nội dung nhỏ, hỗ trợ kỹ thuật.", "1.500.000 - 3.000.000 VNĐ/tháng", "Nên mua sau bàn giao"),
    ]
    t = doc.add_table(rows=len(addons) + 1, cols=4)
    table_style(t)
    for i, h in enumerate(["HẠNG MỤC", "MÔ TẢ", "GIÁ/THỜI GIAN", "ƯU TIÊN"]):
        set_cell_text(t.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(addons, start=1):
        for c, val in enumerate(row):
            set_cell_text(t.cell(r, c), val, size=8.1)

    add_heading(doc, 11, "PHỤ LỤC 3 - TIÊU CHÍ NGHIỆM THU")
    accept = [
        ("01", "Giao diện", "Đủ màn hình chính, đúng nhận diện 4U HOME, hiển thị tốt desktop/mobile.", "□ Đạt"),
        ("02", "Website chính", "Trang chủ, giới thiệu, dịch vụ/báo giá, dự án, chi tiết, danh mục nội thất, blog, liên hệ hoạt động.", "□ Đạt"),
        ("03", "CMS/Admin", "Quản trị được banner, menu, trang tĩnh, dịch vụ, dự án, gallery, bài viết, landing page, form liên hệ/lead và các trường SEO tương ứng.", "□ Đạt"),
        ("04", "Landing/site vệ tinh", "Có template landing page/site vệ tinh và liên kết về trang mẹ theo phạm vi chốt.", "□ Đạt"),
        ("05", "Subdomain", "Cấu hình subdomain theo quyền truy cập Bên A cung cấp, nếu hạng mục này được chốt triển khai.", "□ Đạt"),
        ("06", "SEO cơ bản", "Có metadata, URL thân thiện, sitemap, robots, canonical, schema cơ bản, redirect 301/302 khi cập nhật URL, analytics/tag manager và Meta Pixel/Facebook Pixel nếu Bên A cung cấp mã tracking.", "□ Đạt"),
        ("07", "Thẻ client-side", "Có khu vực/cơ chế cấu hình thẻ HTML/script phía client như mã xác minh, pixel, conversion tag hoặc mã đo lường hợp lệ theo thông tin Bên A cung cấp.", "□ Đạt"),
        ("08", "Form/CTA", "Form liên hệ gửi/ghi nhận được dữ liệu; nút gọi/Zalo/social đúng thông tin cung cấp.", "□ Đạt"),
        ("09", "Responsive/tốc độ", "Không vỡ layout trên mobile/tablet/desktop; tối ưu ảnh/tải trang ở mức cơ bản.", "□ Đạt"),
        ("10", "Bàn giao", "Bàn giao mã nguồn, admin, database/dữ liệu mẫu, hướng dẫn sử dụng cơ bản.", "□ Đạt"),
        ("11", "Ổn định", "Theo dõi và sửa lỗi thuộc phạm vi bảo hành sau nghiệm thu.", "□ Đạt"),
    ]
    t = doc.add_table(rows=len(accept) + 1, cols=4)
    table_style(t)
    for i, h in enumerate(["STT", "HẠNG MỤC", "TIÊU CHÍ NGHIỆM THU", "KQ"]):
        set_cell_text(t.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(accept, start=1):
        for c, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if c in [0, 3] else None
            set_cell_text(t.cell(r, c), val, size=8.0, align=align)

    add_heading(doc, 12, "BIÊN BẢN NGHIỆM THU & BÀN GIAO")
    handover = [
        ("Tên dự án", "Website nội thất 4U HOME"),
        ("Bên A", "................................................................................"),
        ("Bên B", "................................................................................"),
        ("Ngày nghiệm thu", "…… / …… / 2026"),
        ("Tổng giá trị hợp đồng", "23.000.000 VNĐ"),
        ("Hosting/tên miền/subdomain", "................................................................................"),
        ("Kết luận nghiệm thu", "□ Đạt nghiệm thu    □ Cần chỉnh sửa theo ghi chú"),
        ("Ghi chú", "................................................................................"),
    ]
    t = doc.add_table(rows=len(handover), cols=2)
    table_style(t, header=False)
    for r, (a, b) in enumerate(handover):
        set_cell_text(t.cell(r, 0), a, bold=True, size=8.8)
        set_cell_text(t.cell(r, 1), b, size=8.8)

    add_heading(doc, 13, "DANH MỤC HẠNG MỤC BÀN GIAO")
    deliver = [
        "Mã nguồn/project files",
        "Tài khoản super admin/admin",
        "Database/dữ liệu mẫu",
        "Cấu hình hosting/tên miền/subdomain",
        "Template landing page/site vệ tinh",
        "Danh sách redirect 301/302 đã cấu hình nếu có",
        "Cấu hình analytics/tag manager/Meta Pixel nếu có mã tracking",
        "Cấu hình thẻ HTML/script client-side theo phạm vi đã chốt",
        "Tài liệu hướng dẫn sử dụng cơ bản",
        "Danh sách tài khoản/dịch vụ bên thứ ba nếu có",
        "Biên bản lỗi/chỉnh sửa còn lại nếu có",
    ]
    t = doc.add_table(rows=len(deliver) + 1, cols=4)
    table_style(t)
    for i, h in enumerate(["STT", "HẠNG MỤC", "TRẠNG THÁI", "GHI CHÚ"]):
        set_cell_text(t.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r, item in enumerate(deliver, start=1):
        set_cell_text(t.cell(r, 0), f"{r:02d}", size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(t.cell(r, 1), item, size=8.2)
        set_cell_text(t.cell(r, 2), "□ Đã nhận", size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(t.cell(r, 3), "", size=8.2)

    add_heading(doc, 14, "KÝ TÊN XÁC NHẬN")
    add_note(doc, "Hai bên xác nhận đã đọc, hiểu và thống nhất nội dung hợp đồng, phụ lục phạm vi tính năng, phụ lục mua thêm và bảng nghiệm thu bàn giao kèm theo.")
    t = doc.add_table(rows=4, cols=2)
    table_style(t, header=False)
    for c, label in enumerate(["ĐẠI DIỆN BÊN A", "ĐẠI DIỆN BÊN B"]):
        shade(t.cell(0, c), LIGHT)
        set_cell_text(t.cell(0, c), label, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(t.cell(1, 0), "\n\n\n\n", size=10)
    set_cell_text(t.cell(1, 1), "\n\n\n\n", size=10)
    set_cell_text(t.cell(2, 0), "Ký, ghi rõ họ tên", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(t.cell(2, 1), "Ký, ghi rõ họ tên", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(t.cell(3, 0), "…… / …… / 2026", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(t.cell(3, 1), "…… / …… / 2026", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Hợp đồng mẫu 4U HOME - UI/UX & Website - lập ngày 17/05/2026")
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(130, 130, 130)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
